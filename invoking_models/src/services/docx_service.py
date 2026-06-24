import logging
from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from config import settings

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_BYTES = getattr(settings, "DOCX_MAX_FILE_SIZE_BYTES", 25 * 1024 * 1024)  # 25MB


class DocxProcessingError(Exception):
    """Raised for any recoverable DOCX processing failure (bad file, too large, corrupt, etc.)."""


class DocxService:

    # ---------- internal helpers ----------

    @staticmethod
    def _open_docx(file_bytes: bytes) -> DocumentObject:
        if not file_bytes:
            raise DocxProcessingError("Empty file provided.")

        if len(file_bytes) > MAX_FILE_SIZE_BYTES:
            raise DocxProcessingError(
                f"File size {len(file_bytes)} bytes exceeds max allowed "
                f"{MAX_FILE_SIZE_BYTES} bytes."
            )

        try:
            return Document(BytesIO(file_bytes))
        except Exception as exc:
            logger.warning("Failed to open DOCX: %s", exc)
            raise DocxProcessingError("File is not a valid or readable DOCX.") from exc

    @staticmethod
    def _iter_block_items(document: DocumentObject):
        """
        Yields paragraphs and tables in the order they appear in the document body.
        Plain python-docx exposes .paragraphs and .tables separately (losing order);
        this walks the underlying XML body to preserve document order.
        """
        parent_elm = document.element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    @staticmethod
    def _table_to_text(table: Table) -> str:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        return "\n".join(rows_text)

    # ---------- public API ----------

    @staticmethod
    def extract_text(file_bytes: bytes) -> str:
        """
        Extracts paragraph and table text in document order.
        Note: tables nested inside table cells are not separately recursed —
        only the cell's own paragraph text is captured (python-docx limitation).
        Headers/footers are not included.
        """
        document = DocxService._open_docx(file_bytes)

        parts = []
        paragraph_count = 0
        table_count = 0

        for block in DocxService._iter_block_items(document):
            try:
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    if text:
                        parts.append(text)
                        paragraph_count += 1
                elif isinstance(block, Table):
                    table_text = DocxService._table_to_text(block)
                    if table_text.strip():
                        parts.append(table_text)
                        table_count += 1
            except Exception as exc:
                logger.warning("Failed to extract a block from DOCX: %s", exc)
                continue

        result = "\n".join(parts).strip()

        logger.info(
            "DOCX extracted: paragraphs=%s tables=%s total_chars=%s",
            paragraph_count, table_count, len(result),
        )

        if not result:
            logger.warning("DOCX extraction produced no text content.")

        return result